"""
Word Builder - Tạo báo cáo tóm tắt tình hình dịch bệnh dạng .docx.
"""
import io
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, hex_color: str):
    """Đặt màu nền cho cell trong bảng Word."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_paragraph(doc: Document, text: str, bold: bool = False, size: int = 12,
                   color: Optional[str] = None, align=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before: int = 0, space_after: int = 6) -> None:
    """Thêm đoạn văn với định dạng chuẩn."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def build_word_report(report_data: dict) -> bytes:
    """
    Tạo báo cáo Word tóm tắt dịch bệnh.
    
    Args:
        report_data: Dict từ generator.get_report_data()
    
    Returns:
        bytes của file .docx
    """
    doc = Document()

    # Thiết lập margin trang
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    generated_at: datetime = report_data.get("generated_at", datetime.utcnow())
    scope_hours: int = report_data.get("scope_hours", 72)
    start_date: datetime = report_data.get("start_date", generated_at)
    end_date: datetime = report_data.get("end_date", generated_at)
    overview: dict = report_data.get("overview", {})
    top_events = report_data.get("top_events", [])
    top_diseases = report_data.get("top_diseases", [])
    alert_articles = report_data.get("alert_articles", [])

    # =========================================================
    # TIÊU ĐỀ
    # =========================================================
    _add_paragraph(
        doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_paragraph(
        doc, "Độc lập - Tự do - Hạnh phúc",
        bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
    )
    _add_paragraph(doc, "————————", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    _add_paragraph(
        doc,
        "BÁO CÁO GIÁM SÁT DỊCH BỆNH DỰA VÀO SỰ KIỆN (EBS)",
        bold=True, size=14, color="1565C0",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4,
    )

    _add_paragraph(
        doc,
        f"Kỳ báo cáo: {start_date.strftime('%d/%m/%Y %H:%M')} — {end_date.strftime('%d/%m/%Y %H:%M')} "
        f"(trong {scope_hours} giờ qua)",
        size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
    )

    # =========================================================
    # I. TÓM TẮT THỐNG KÊ
    # =========================================================
    _add_paragraph(doc, "I. TÓM TẮT THỐNG KÊ CHÍNH", bold=True, size=12, color="1565C0", space_before=8)

    # Bảng tổng quan
    stats_table = doc.add_table(rows=2, cols=3)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats_table.style = "Table Grid"

    headers = ["Tổng số bài báo", "Tổng số ca bệnh", "Tín hiệu cảnh báo"]
    values = [
        str(overview.get("total_articles", 0)),
        str(overview.get("total_cases", 0)),
        str(overview.get("alert_count", 0)),
    ]
    for i, (h, v) in enumerate(zip(headers, values)):
        hc = stats_table.cell(0, i)
        hc.text = h
        hc.paragraphs[0].runs[0].bold = True
        hc.paragraphs[0].runs[0].font.size = Pt(10)
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(hc, "1976D2")
        for run in hc.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

        vc = stats_table.cell(1, i)
        vc.text = v
        vc.paragraphs[0].runs[0].font.size = Pt(16)
        vc.paragraphs[0].runs[0].bold = True
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================================================
    # II. TOP SỰ KIỆN NỔI BẬT
    # =========================================================
    _add_paragraph(doc, "II. CÁC SỰ KIỆN DỊCH BỆNH NỔI BẬT", bold=True, size=12, color="1565C0", space_before=14)

    if top_events:
        for idx, event in enumerate(top_events[:5], start=1):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            run_title = p.add_run(f"{event.canonical_title}")
            run_title.bold = True
            run_title.font.size = Pt(11)
            run_title.font.name = "Times New Roman"

            if event.location or event.case_count:
                detail_parts = []
                if event.location:
                    detail_parts.append(f"Địa điểm: {event.location}")
                if event.case_count:
                    detail_parts.append(f"Số ca: {event.case_count}")
                if event.article_count:
                    detail_parts.append(f"Được đưa tin bởi {event.article_count} bài báo")
                detail_p = doc.add_paragraph("   " + " | ".join(detail_parts))
                detail_p.paragraph_format.space_before = Pt(0)
                detail_p.paragraph_format.space_after = Pt(4)
                for run in detail_p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(97, 97, 97)
                    run.font.name = "Times New Roman"
    else:
        _add_paragraph(doc, "Không có sự kiện nổi bật trong khoảng thời gian này.", size=11, color="757575")

    # =========================================================
    # III. TOP DỊCH BỆNH ĐƯỢC NHẮC ĐẾN NHIỀU NHẤT
    # =========================================================
    _add_paragraph(doc, "III. TOP DỊCH BỆNH ĐƯỢC NHẮC ĐẾN NHIỀU NHẤT", bold=True, size=12, color="1565C0", space_before=14)

    if top_diseases:
        dis_table = doc.add_table(rows=len(top_diseases[:8]) + 1, cols=2)
        dis_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        dis_table.style = "Table Grid"

        h0 = dis_table.cell(0, 0)
        h0.text = "Dịch bệnh / Từ khóa"
        h0.paragraphs[0].runs[0].bold = True
        h0.paragraphs[0].runs[0].font.size = Pt(10)
        _set_cell_bg(h0, "1976D2")
        for run in h0.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

        h1 = dis_table.cell(0, 1)
        h1.text = "Số bài báo"
        h1.paragraphs[0].runs[0].bold = True
        h1.paragraphs[0].runs[0].font.size = Pt(10)
        h1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(h1, "1976D2")
        for run in h1.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

        for i, disease in enumerate(top_diseases[:8], start=1):
            c0 = dis_table.cell(i, 0)
            c0.text = disease.get("disease_name", "")
            c0.paragraphs[0].runs[0].font.size = Pt(10)

            c1 = dis_table.cell(i, 1)
            c1.text = str(disease.get("article_count", 0))
            c1.paragraphs[0].runs[0].font.size = Pt(10)
            c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        _add_paragraph(doc, "Chưa có dữ liệu.", size=11, color="757575")

    # =========================================================
    # IV. TÍN HIỆU CẢNH BÁO CẦN CHÚ Ý
    # =========================================================
    _add_paragraph(doc, "IV. TÍN HIỆU CẢNH BÁO CẦN CHÚ Ý", bold=True, size=12, color="C62828", space_before=14)

    if alert_articles:
        for article in alert_articles:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            run_t = p.add_run(article.title or "Không có tiêu đề")
            run_t.font.size = Pt(11)
            run_t.font.name = "Times New Roman"
            run_t.font.color.rgb = RGBColor(198, 40, 40)

            if article.source or article.published_date:
                detail_parts = []
                if article.source:
                    detail_parts.append(f"Nguồn: {article.source}")
                if article.published_date:
                    detail_parts.append(f"Ngày: {article.published_date.strftime('%d/%m/%Y')}")
                detail_p = doc.add_paragraph("   " + " | ".join(detail_parts))
                detail_p.paragraph_format.space_before = Pt(0)
                detail_p.paragraph_format.space_after = Pt(4)
                for run in detail_p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(97, 97, 97)
                    run.font.name = "Times New Roman"
    else:
        _add_paragraph(doc, "Không phát hiện tín hiệu cảnh báo đặc biệt.", size=11, color="2E7D32")

    # =========================================================
    # CHỮ KÝ
    # =========================================================
    _add_paragraph(doc, "", space_before=20)
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sig = sig_p.add_run(
        f"Hà Nội, ngày {generated_at.day} tháng {generated_at.month} năm {generated_at.year}\n"
        "Người lập báo cáo\n\n\n\n"
        "(Ký, ghi rõ họ tên)"
    )
    run_sig.font.size = Pt(11)
    run_sig.font.name = "Times New Roman"

    # GHI CHÚ CUỐI
    _add_paragraph(
        doc,
        f"* Báo cáo được tạo tự động bởi Hệ thống Epi Scout AI lúc {generated_at.strftime('%H:%M %d/%m/%Y')}. "
        "Cán bộ y tế cần xem xét và xác minh trước khi sử dụng chính thức.",
        size=9, color="757575", space_before=12,
    )

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
